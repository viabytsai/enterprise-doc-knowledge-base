from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path


@dataclass
class Page:
    number: int | None
    text: str
    section_title: str | None = None


@dataclass
class ParsedDocument:
    pages: list[Page]

    @property
    def page_count(self) -> int:
        return len(self.pages)


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix in {".md", ".txt"}:
        return _parse_text(path)
    raise ValueError(f"不支持的文件类型：{suffix}")


def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF 解析依赖尚未安装") from exc

    reader = PdfReader(str(path))
    pages = [
        Page(number=index, text=_clean_text(page.extract_text() or ""))
        for index, page in enumerate(reader.pages, start=1)
    ]
    return ParsedDocument(pages=[page for page in pages if page.text])


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX 解析依赖尚未安装") from exc

    document = Document(path)
    pages: list[Page] = []
    current_title: str | None = None
    paragraphs: list[str] = []

    def flush() -> None:
        if paragraphs:
            pages.append(
                Page(
                    number=None,
                    text=_clean_text("\n".join(paragraphs)),
                    section_title=current_title,
                )
            )
            paragraphs.clear()

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            flush()
            current_title = text
        else:
            paragraphs.append(text)
    flush()
    return ParsedDocument(pages=pages or [Page(number=None, text="")])


def _parse_text(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    if path.suffix.lower() == ".md":
        pages: list[Page] = []
        title: str | None = None
        body: list[str] = []
        for line in text.splitlines():
            if re.match(r"^#{1,6}\s+", line):
                if body:
                    pages.append(Page(number=None, text=_clean_text("\n".join(body)), section_title=title))
                    body = []
                title = re.sub(r"^#{1,6}\s+", "", line).strip()
            else:
                body.append(line)
        if body:
            pages.append(Page(number=None, text=_clean_text("\n".join(body)), section_title=title))
        return ParsedDocument(pages=[page for page in pages if page.text])
    return ParsedDocument(pages=[Page(number=None, text=_clean_text(text))])


def _clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_pages(
    parsed: ParsedDocument, chunk_size: int, overlap: int
) -> list[dict[str, object]]:
    chunks: list[dict[str, object]] = []
    for page in parsed.pages:
        text = page.text.strip()
        if not text:
            continue
        start = 0
        while start < len(text):
            target_end = min(start + chunk_size, len(text))
            end = _find_boundary(text, start, target_end)
            content = text[start:end].strip()
            if content:
                chunks.append(
                    {
                        "content": content,
                        "page_number": page.number,
                        "section_title": page.section_title,
                    }
                )
            if end >= len(text):
                break
            start = max(end - overlap, start + 1)
    return chunks


def _find_boundary(text: str, start: int, target_end: int) -> int:
    if target_end >= len(text):
        return len(text)
    minimum = start + int((target_end - start) * 0.65)
    for separator in ("\n\n", "\n", "。", "；", "！", "？", ". "):
        index = text.rfind(separator, minimum, target_end)
        if index != -1:
            return index + len(separator)
    return target_end

